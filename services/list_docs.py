from docx import Document


def replace_text_in_docx(doc_path: str, index: int, new_text: str) -> bool:
    """Замена обычного текста в docx по индексу"""

    doc = Document(doc_path)

    done_if = False
    if index < len(doc.paragraphs):
        para = doc.paragraphs[index]
        inline = para.runs
        for i in range(len(inline)):
            if i > 0:
                inline[i].text = ''
            else:
                inline[0].text = new_text

        if index == len(doc.paragraphs) - 1:
            done_if = True

    doc.save(doc_path)
    return done_if


def replace_header_footer_text(doc_path: str, type_col: str, index: int, new_text) -> bool:
    """Замена текста в колонтитулах docx по индексу"""
    doc = Document(doc_path)

    done_if = False
    for section in doc.sections:

        match type_col:
            case 'up':
                colont = section.header
            case _:
                colont = section.footer

        if index < len(colont.paragraphs):
            para = colont.paragraphs[index]
            inline = para.runs
            for i in range(len(inline)):
                if i > 0:
                    inline[i].text = ''
                else:
                    inline[0].text = new_text

        if index == len(colont.paragraphs) - 1:
            done_if = True

    doc.save(doc_path)
    return done_if


# Таблица
# Функция для замены текста в таблицах
def replace_text_in_table(doc_path: str, table_index: int, row_index, cell_index, n_index: int, n_row_index: int,
                          n_cell_index: int, index: int, new_text):
    doc = Document(doc_path)

    done_if = False

    if doc.tables:

        table = doc.tables[table_index]

        row = table.rows[0]

        if cell_index > len(row.cells) - 1:
            cell_index = len(row.cells) - 1

        cell = table.cell(row_index, cell_index)


        # Проверяем, является ли это последней таблицей
        is_last_table = table_index == len(doc.tables) - 1
        # Проверяем, является ли это последней ячейкой в таблице
        is_last_cell = (row_index == len(table.rows) - 1) and (cell_index == len(table.columns) - 1)

        if cell.tables and not n_index == -1:
            n_table = cell.tables[n_index]
            n_cell = n_table.cell(n_row_index, n_cell_index)

            # Проверяем, является ли это последней таблицей
            is_last_n_table = n_index == len(cell.tables) - 1
            # Проверяем, является ли это последней ячейкой в таблице
            is_last_n_cell = (n_row_index == len(n_table.rows) - 1) and (n_cell_index == len(n_table.columns) - 1)
            is_last_n_parag = index == len(n_cell.paragraphs) - 1

            if index < len(n_cell.paragraphs):
                para = n_cell.paragraphs[index]
                inline = para.runs

                for i in range(len(inline)):
                    if i > 0:
                        inline[i].text = ''
                    else:
                        inline[0].text = new_text

        else:
            is_last_parag = index == len(cell.paragraphs) - 1
            if index < len(cell.paragraphs):
                para = cell.paragraphs[index]
                inline = para.runs

                for i in range(len(inline)):
                    if i > 0:
                        inline[i].text = ''
                    else:
                        inline[0].text = new_text

        doc.save(doc_path)

        #if is_last_table and is_last_n_table and is_last_cell and is_last_parag and is_last_n_parag and is_last_n_cell:
        #    done_if = True
        return done_if
