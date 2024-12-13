import os
import subprocess
from abc import ABC, abstractmethod

from docx import Document
from core import config
from requests.compat import chardet

settings = config.get_settings()


class FileReader(ABC):
    def __init__(self, method):
        self.method = method

    @abstractmethod
    def file_read(self, file: str):
        pass


class DocxReader(FileReader):

    def file_read(self, file_name: str):
        docx = Document(file_name)

        match self.method:
            case 1:
                docs_str = ''

                for section in docx.sections:

                    for paragraph in section.header.paragraphs:
                        docs_str += paragraph.text + '\n'

                    for paragraph in section.footer.paragraphs:
                        docs_str += paragraph.text + '\n'

                for i in range(len(docx.paragraphs)):
                    # if i <= len(docx.paragraphs) - 1:
                    if len(docx.tables) > 0:
                        docs_str += docx.paragraphs[i].text + '\n'
                    else:
                        if i < len(docx.paragraphs) - 1:
                            docs_str += docx.paragraphs[i].text + '\n'
                        else:
                            docs_str += docx.paragraphs[i].text

                for table in docx.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            docs_str += cell.text + '\n'

                            if cell.tables:
                                for nested_table in cell.tables:
                                    for nested_row in nested_table.rows:
                                        for nested_cell in nested_row.cells:
                                            docs_str += "  " + nested_cell.text + '\n'  # Добавляем отступ для

                return docs_str
            case _:
                list_docs_str = []
                for paragraph in docx.paragraphs:
                    list_docs_str.append(paragraph)

                return list_docs_str

    def colontituls_read(self, file_name: str):
        docx = Document(file_name)

        headers = []
        footers = []

        for section in docx.sections:

            for paragraph in section.header.paragraphs:
                headers.append(paragraph.text)

            for paragraph in section.footer.paragraphs:
                footers.append(paragraph.text)

        return headers, footers

    def table_read(self, file_name: str):
        docx = Document(file_name)

        list_table_data = []
        list_n_table_data = []

        for i_table in range(len(docx.tables)):
            for i_row in range(len(docx.tables[i_table].rows)):
                for i_cell in range(len(docx.tables[i_table].rows[i_row].cells)):
                    text_cell = docx.tables[i_table].rows[i_row].cells[i_cell].text
                    t_cell = text_cell.split('\n')
                    for i in range(len(t_cell)):
                        cur_cell = t_cell[i]
                        if cur_cell == '':
                            cur_cell = ' '
                        list_table_data.append(
                            f'table:{i_table}row:{i_row}cell:{i_cell}n_table:-1n_row:-1n_cell:-1index:{i}text:{cur_cell}')

                    if docx.tables[i_table].rows[i_row].cells[i_cell].tables:
                        for i_nested_table, nested_table in enumerate(
                                docx.tables[i_table].rows[i_row].cells[i_cell].tables):
                            for i_nested_row, nested_row in enumerate(nested_table.rows):
                                for i_nested_cell, nested_cell in enumerate(nested_row.cells):
                                    n_cell = nested_cell.text.split('\n')
                                    for i in range(len(n_cell)):
                                        n_cur_cell = n_cell[i]
                                        if n_cur_cell == '':
                                            n_cur_cell = ' '
                                        list_n_table_data.append(
                                            f'table:{i_table}row:{i_row}cell:{i_cell}n_table:{i_nested_table}n_row:{i_nested_row}n_cell:{i_nested_cell}index:{i}text:{n_cur_cell}'
                                        )

        return list_table_data, list_n_table_data


class TXTReader(FileReader):

    def file_read(self, file: str):
        with open(file, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding']

        with open(file, 'r', encoding=encoding) as f_read:
            lines = f_read.readlines()
        match self.method:
            case 1:
                data_str = ''
                for data in lines:
                    data_str += data
                return data_str
            case _:
                return lines


class PDFReader(FileReader):
    def file_read(self, file: str):
        try:
            if not os.path.isdir('./temp'):
                os.mkdir('./temp')
            # Запускаем команду pdftotext
            result = subprocess.run(['pdftotext', '-enc', 'UTF-8', file, f'./temp/{file[file.rfind('\\'):-3]}txt'],
                                    check=True, text=True, capture_output=True)
            # print(result.stdout)  # Выводим извлеченный текст
            return TXTReader(method=1).file_read(f'./temp/{file[file.rfind('\\'):-3]}txt')

        except subprocess.CalledProcessError as e:
            print(f"Ошибка при извлечении текста: {e}")
        except FileNotFoundError:
            print("Не удалось найти исполняемый файл pdftotext. Убедитесь, что xpdf установлен и путь добавлен в PATH.")
